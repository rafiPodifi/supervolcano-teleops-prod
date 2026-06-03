#!/usr/bin/env python3
"""Generate docs/ARCHITECTURE.docx — executive technical architecture overview.

One-off generator built on python-docx. Run:
    python3 scripts/gen-architecture-docx.py
Produces docs/ARCHITECTURE.docx (stakeholder-altitude, ~8-12 pages).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "docs", "ARCHITECTURE.docx")

ACCENT = RGBColor(0xC0, 0x39, 0x2B)   # volcano red
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def _color_heading(style_name, color, size):
    st = doc.styles[style_name]
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.name = "Calibri"


_color_heading("Title", DARK, 30)
_color_heading("Heading 1", ACCENT, 18)
_color_heading("Heading 2", DARK, 14)
_color_heading("Heading 3", GREY, 11.5)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text=""):
    p = doc.add_paragraph(text)
    return p


def lead(text):
    """Slightly larger intro paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.color.rgb = GREY
    return p


def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            label, rest = it
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(label + " — ")
            r.bold = True
            p.add_run(rest)
        else:
            doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def page_break():
    doc.add_page_break()


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light List Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "C0392B")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# =====================================================================
# TITLE PAGE
# =====================================================================
for _ in range(6):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("SuperVolcano Teleops")
r.font.size = Pt(34)
r.bold = True
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Technical Architecture Overview")
r.font.size = Pt(18)
r.font.color.rgb = DARK

for _ in range(2):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Platform for robot-training video collection\nWeb dashboard · Android mobile app · Google Cloud backend")
r.font.size = Pt(11.5)
r.font.color.rgb = GREY

for _ in range(8):
    doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Version 1.0  ·  June 4, 2026  ·  Confidential")
r.font.size = Pt(10)
r.font.color.rgb = GREY
page_break()

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
h1("1. Executive Summary")
lead(
    "SuperVolcano is a platform that turns everyday cleaning and teleoperation "
    "work into a curated, anonymized video corpus used to train robots."
)
para(
    "Field workers (cleaners and OEM teleoperators) record video of real tasks "
    "through an Android mobile app. That footage flows into a Google Cloud backend "
    "where it is automatically analyzed, classified, reviewed by internal staff, and — "
    "once approved — promoted into an anonymized training dataset that robotics "
    "partners consume through a dedicated API."
)
para("The system has three software components, all in a single repository (monorepo):")
bullets([
    ("Web dashboard", "Next.js 14 application for internal staff, customer "
     "organizations, and the robot-facing data API."),
    ("Mobile app", "Expo / React Native Android app for field workers — the data "
     "capture front line, built to keep working offline."),
    ("Cloud functions & backend services", "Firebase Cloud Functions plus backend "
     "services that run the video-intelligence pipeline and external integrations."),
])
para(
    "The platform serves two business models simultaneously: a B2B model for OEM "
    "robotics testing, and a B2C model for property-management cleaning. A six-role "
    "permission system keeps the two worlds — and each customer's data — isolated."
)

# =====================================================================
# SYSTEM OVERVIEW
# =====================================================================
h1("2. System Overview")
para(
    "The repository is organized as a monorepo with three top-level components, "
    "each deployable independently:"
)
table(
    ["Component", "Location", "Technology", "Purpose"],
    [
        ["Web dashboard", "/src", "Next.js 14 (App Router), TypeScript, Tailwind",
         "Admin, customer, and robot-API portals"],
        ["Mobile app", "/mobile-app", "Expo / React Native (Android)",
         "Field-worker video capture"],
        ["Cloud functions", "/functions", "Firebase Cloud Functions (Node)",
         "Scheduled sync + backend jobs"],
        ["Infrastructure", "/infra/terraform", "Terraform (GCP)",
         "All cloud resources as code"],
    ],
    widths=[1.3, 1.5, 2.2, 2.2],
)

h2("Technology Stack")
table(
    ["Layer", "Technology"],
    [
        ["Web framework", "Next.js 14 (App Router) + TypeScript"],
        ["UI", "TailwindCSS + shadcn/ui"],
        ["Mobile", "Expo ~54 / React Native 0.81 (Android), custom native modules"],
        ["Operational database", "Cloud Firestore (named databases per environment)"],
        ["Analytical / robot database", "PostgreSQL (Cloud SQL)"],
        ["Authentication", "Firebase Identity Platform (multi-tenant)"],
        ["Video AI", "Google Cloud Video Intelligence API"],
        ["Hosting", "Google Cloud Run (web), Cloud Functions (jobs)"],
        ["Infrastructure as code", "Terraform"],
        ["CI / CD", "GitHub Actions + Workload Identity Federation"],
    ],
    widths=[2.5, 4.5],
)

h2("High-Level Flow")
para("At a glance, data moves through the platform like this:")
para(
    "Field worker records video (mobile app)  →  video uploaded to Firebase Storage  "
    "→  metadata written to Firestore  →  video-intelligence pipeline analyzes and "
    "classifies it  →  internal staff review and approve  →  anonymized record copied "
    "to the PostgreSQL training corpus  →  robotics partners read it via the Robot API."
)
page_break()

# =====================================================================
# BUSINESS MODEL & ROLES
# =====================================================================
h1("3. Business Model & Roles")
para("The platform supports two distinct business models on shared infrastructure:")
bullets([
    ("B2B — OEM Robotics Testing", "Robotics manufacturers (OEM partners) whose "
     "teleoperators record task footage at assigned locations."),
    ("B2C — Property Management", "Property owners who manage their own properties "
     "and dispatch cleaners to record cleaning work."),
])

h2("Roles")
para(
    "Access is governed by six roles across the two models. Permission enforcement is "
    "centralized in the web backend (src/lib/auth/permissions.ts) and mirrored in the "
    "database security rules."
)
table(
    ["Role", "Model", "Organization", "Access", "Platform"],
    [
        ["admin / superadmin", "Internal", "sv:internal", "Full platform management", "Web"],
        ["partner_manager", "B2B", "oem:<slug>", "Assigned locations; cannot create locations", "Web + Mobile"],
        ["oem_teleoperator", "B2B", "oem:<slug>", "Record at assigned locations", "Mobile only"],
        ["location_owner", "B2C", "owner:<slug>", "Creates and owns own properties", "Web + Mobile"],
        ["location_cleaner", "B2C", "owner:<slug>", "Record at assigned properties", "Mobile only"],
    ],
    widths=[1.5, 0.9, 1.3, 2.0, 1.3],
)
para(
    "Every data-returning endpoint filters results by the caller's role and "
    "organization, so each customer sees only their own locations, tasks, and media."
)

# =====================================================================
# WEB DASHBOARD
# =====================================================================
h1("4. Web Dashboard")
para(
    "The Next.js application hosts three logically distinct portals within one "
    "deployment:"
)
bullets([
    ("/admin", "SuperVolcano internal team — full management of organizations, "
     "locations, users, tasks, and the robot-intelligence/training library."),
    ("/org", "Customer organizations — a shared portal where managers and field "
     "workers see role-filtered views of their locations and tasks."),
    ("/api/robot/v1/*", "Robot Intelligence API — the read path for robotics "
     "partners, backed exclusively by PostgreSQL."),
])

h2("Layered Design")
para("The web codebase separates concerns into clear layers:")
bullets([
    ("API routes", "Verify the caller's token, enforce permissions, then delegate "
     "to the data layer."),
    ("Repositories", "All Firestore reads and writes for human-facing data flow "
     "through a single repository layer (src/lib/repositories)."),
    ("Validation", "All API input is validated with Zod schemas before it reaches "
     "the database."),
    ("Firebase Admin", "Server-side database access uses the Firebase Admin SDK, "
     "which runs with elevated privileges and is governed by application code rather "
     "than client security rules."),
])

# =====================================================================
# MOBILE APP
# =====================================================================
h1("5. Mobile App")
lead(
    "The Android app is the data-capture front line. It is built to keep working "
    "when the network does not."
)
h2("Role-Based Navigation")
para(
    "On sign-in the app routes the user into one of three experiences based on their "
    "role: a cleaner/teleoperator recording flow, a property-owner management flow, "
    "or a member flow. Cleaners land directly on the camera screen — recording is "
    "one tap away."
)

h2("Camera-First Capture with GPS Auto-Assign")
para(
    "When a cleaner opens the app, it requests a GPS fix and automatically binds the "
    "recording to the nearest location they are assigned to — no manual selection "
    "needed. A manual location picker is always available as an override, and footage "
    "can also be recorded generically and assigned to a job later."
)

h2("External USB Camera")
para(
    "A custom Android native module supports external USB (UVC) cameras in addition "
    "to the built-in camera, negotiating video format and quality and handling "
    "USB attach/detach events. This lets field workers capture higher-fidelity "
    "footage with dedicated hardware."
)

h2("Resilient Upload Queue")
para(
    "Video is uploaded directly from the device to Firebase Storage using a resumable "
    "protocol, then its metadata is posted to the backend. A persistent upload queue "
    "is the backbone of the app's reliability:"
)
bullets([
    "Recordings are segmented and queued locally, surviving app restarts and device reboots.",
    "Uploads run in the background and retry automatically with increasing back-off delays.",
    "The queue respects online/offline transitions and keeps the device awake only while actively uploading.",
    "Each cleaner's recorded hours are attributed on the backend from the uploaded metadata.",
])
page_break()

# =====================================================================
# DATA ARCHITECTURE
# =====================================================================
h1("6. Data Architecture")
lead("The platform uses two databases with a strict, never-mixed boundary.")
table(
    ["", "Firestore", "PostgreSQL"],
    [
        ["Role", "Source of truth for all human-facing data",
         "Curated robot training corpus + AI processing queue"],
        ["Consumers", "Admin, org, and mobile endpoints", "Robot API only (/api/robot/v1/*)"],
        ["Access", "Firebase Admin SDK (adminDb)", "Direct SQL (Cloud SQL)"],
        ["Writes", "All application writes", "Only the video-intelligence pipeline"],
    ],
    widths=[1.1, 2.9, 2.9],
)
para(
    "The two databases are never queried together: admin and org endpoints touch "
    "Firestore only; the robot API touches PostgreSQL only. Data moves one way — "
    "from Firestore into the PostgreSQL corpus — and only when a human approves a "
    "video for training. There is no automatic two-way sync."
)
h2("What Lives Where")
bullets([
    ("Firestore", "Organizations, users, locations and their structure (floors / "
     "rooms / targets / actions), tasks, sessions, and media metadata including AI "
     "annotations."),
    ("PostgreSQL", "Anonymized training videos, the video-processing queue, robot "
     "API keys, and API usage metering."),
])

# =====================================================================
# VIDEO INTELLIGENCE
# =====================================================================
h1("7. Video Intelligence Pipeline")
para(
    "Approved footage becomes robot-training data through an automated, "
    "queue-driven pipeline:"
)
numbered([
    "Upload — a recording lands in Firebase Storage and its metadata is written to Firestore.",
    "Queue — the video is added to a durable processing queue in PostgreSQL with a priority.",
    "Analyze — a worker sends the video to Google Cloud Video Intelligence, which returns labels, objects, and shot segments.",
    "Classify — the pipeline filters the raw labels and derives a room type (kitchen, bathroom, etc.), action types (cleaning, organizing, …), a quality score, and duration.",
    "Review — internal staff review the classified video in the dashboard and approve or reject it for training.",
    "Promote — on approval, an anonymized record (room type, action types, object labels, quality — no location or user identity) is copied to the PostgreSQL training_videos corpus.",
])
para(
    "The queue is built for durability: it survives crashes, processes work without "
    "lock contention, and retries failed items up to a fixed limit before flagging "
    "them."
)

# =====================================================================
# AUTH & SECURITY
# =====================================================================
h1("8. Authentication & Security")
h2("Multi-Tenant Identity")
para(
    "Authentication runs on Firebase Identity Platform in multi-tenant mode. Users "
    "live inside a per-environment tenant user pool (staging or production), not the "
    "project-level pool. Every user-management operation is scoped to the caller's "
    "tenant; only token verification is tenant-agnostic."
)
h2("Dual-Stored Identity")
para(
    "Each user exists in two places that the platform keeps in sync: the Identity "
    "Platform record (which carries role and organization as custom claims on the "
    "login token) and a Firestore user profile document. The admin user editor "
    "explicitly detects and resolves drift between the two."
)
h2("Security Rules & Secrets")
bullets([
    ("Firestore & Storage rules", "Database and file access is gated by security "
     "rules that evaluate the user's role and organization claims, enforcing tenant "
     "isolation at the data layer."),
    ("Secret management", "Credentials and API keys are stored in Google Secret "
     "Manager, scoped per environment so staging credentials can never reach "
     "production resources."),
    ("Least privilege", "Each running service uses a dedicated service account "
     "limited to exactly the resources it needs."),
])
page_break()

# =====================================================================
# INFRASTRUCTURE & DEPLOYMENT
# =====================================================================
h1("9. Infrastructure & Deployment")
para(
    "All cloud infrastructure runs on a single Google Cloud project hosting two "
    "fully isolated environments. Every resource is defined as code in Terraform — "
    "nothing is created by hand."
)
table(
    ["Concern", "Staging", "Production"],
    [
        ["Web service (Cloud Run)", "supervolcano-web-staging", "supervolcano-web-prod"],
        ["Firestore database", "staging-db", "prod-db"],
        ["Cloud SQL", "sv-sql-staging (zonal)", "sv-sql-prod (regional HA, backups)"],
        ["Identity tenant", "staging-*", "prod-*"],
        ["Deploy trigger", "Push to main", "Git tag v* (manual approval)"],
    ],
    widths=[1.9, 2.4, 2.7],
)

h2("Network & Data Isolation")
bullets([
    "Cloud SQL has no public IP — it is reachable only privately from Cloud Run over a VPC connector.",
    "Production is publicly reachable; staging requires authentication.",
    "Production adds high availability, point-in-time recovery, backups, and deletion protection.",
])

h2("CI / CD")
para("Continuous integration and deployment are fully automated through GitHub Actions:")
bullets([
    ("On every pull request", "Type-checking and linting run on the web, mobile, "
     "and functions code."),
    ("On push to main", "The web app is built into a container, pushed, and deployed "
     "to the staging Cloud Run service, followed by a security-rules deploy."),
    ("On a version tag (v*)", "The same flow targets production, gated behind a "
     "manual approval step."),
    ("Authentication", "CI authenticates to Google Cloud keylessly via Workload "
     "Identity Federation — no long-lived service-account keys exist."),
])
para(
    "Note: browser-facing configuration values are baked into the web app at build "
    "time, so changing them requires a rebuild rather than a settings change."
)

# =====================================================================
# PRINCIPLES
# =====================================================================
h1("10. Key Architectural Principles")
bullets([
    ("Source-of-truth separation", "Firestore owns human data; PostgreSQL owns the "
     "robot corpus. One-way, approval-driven promotion between them."),
    ("Role-based access control", "Six roles, centrally enforced in application code "
     "and again in database security rules."),
    ("Tenant isolation", "Separate identity pools, databases, and secrets per "
     "environment; customer data scoped by organization."),
    ("Durable, queue-based processing", "Video analysis survives failures and retries "
     "automatically."),
    ("Offline-first mobile", "Capture and upload keep working without a network and "
     "recover on their own."),
    ("Everything as code", "Infrastructure in Terraform, deploys in GitHub Actions, "
     "keyless cloud authentication."),
])

# =====================================================================
# APPENDIX
# =====================================================================
h1("Appendix A — Glossary")
table(
    ["Term", "Meaning"],
    [
        ["OEM", "Original Equipment Manufacturer — a robotics partner (B2B)."],
        ["Teleoperator", "A field worker recording task footage for an OEM partner."],
        ["Training corpus", "The anonymized PostgreSQL dataset robots learn from."],
        ["UVC", "USB Video Class — standard for external USB cameras."],
        ["Tenant", "An isolated user pool in Identity Platform (per environment)."],
        ["WIF", "Workload Identity Federation — keyless CI authentication to GCP."],
        ["Cloud Run", "Google's managed container hosting (runs the web app)."],
    ],
    widths=[1.6, 5.4],
)

h1("Appendix B — Further Reading (in repo)")
para("This document summarizes detailed material maintained inside the repository:")
bullets([
    ("docs/architecture/ARCHITECTURE.md", "Full architecture reference: roles, data "
     "model, API patterns, workflows."),
    ("docs/architecture/data-layer.md", "The Firestore vs PostgreSQL boundary rules."),
    ("docs/architecture/ROBOT_API.md", "Robot Intelligence API endpoints."),
    ("docs/architecture/VIDEO_INTELLIGENCE_IMPLEMENTATION.md", "Pipeline internals."),
    ("CLAUDE.md", "Developer guidance, deployment details, mobile specifics."),
    ("infra/terraform/README.md", "Infrastructure setup and first-run steps."),
])

doc.save(OUT)
print("Wrote", OUT)
