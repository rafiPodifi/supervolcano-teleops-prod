#!/usr/bin/env python3
"""Generate customer-facing infrastructure options docx covering Cloud SQL,
Cloud Run, networking, storage, and bundled cost scenarios.

Run:
    /Users/ahmadrafi/miniconda3/bin/python3 \
        scripts/generate-infra-options-docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HEADER_BLUE = "1F3864"
ROW_ALT = "F2F2F2"
ROW_HIGHLIGHT = "E2EFDA"  # soft green for recommended row
GREY = RGBColor(0x77, 0x77, 0x77)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        run = p.add_run(bold_lead)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(" — " + text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, highlight_idx=None):
    """Create a styled table. highlight_idx = body row index (0-based) to tint."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], HEADER_BLUE)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if highlight_idx is not None and r == highlight_idx:
                shade_cell(cell, ROW_HIGHLIGHT)
            elif r % 2 == 1:
                shade_cell(cell, ROW_ALT)
    return table


def add_option_block(doc, title, intro, pros, cons):
    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    add_para(doc, intro, size=11)
    p = doc.add_paragraph()
    run = p.add_run("Pros")
    run.bold = True
    run.font.size = Pt(10)
    for item in pros:
        add_bullet(doc, item)
    p = doc.add_paragraph()
    run = p.add_run("Cons")
    run.bold = True
    run.font.size = Pt(10)
    for item in cons:
        add_bullet(doc, item)
    doc.add_paragraph()


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ─── Cover ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUPERVOLCANO TELEOPS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Infrastructure Options & Cost Plan")
    run.bold = True
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GCP migration · single project · staging + production")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Podifi · 2026-05-09 · Confidential — Internal Use Only")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    doc.add_paragraph()

    # ─── Executive summary ───────────────────────────────────
    add_heading(doc, "Executive summary", level=1)
    add_para(
        doc,
        "The Supervolcano platform is moving from Vercel + Firebase + Neon to Google "
        "Cloud Platform under a single client-owned project. Both staging and production "
        "environments live in that one project, isolated by Identity Platform tenants, "
        "Firestore named databases, separate Cloud SQL instances, and per-environment "
        "service accounts and storage buckets. The web application runs as a container "
        "on Cloud Run; the mobile app continues to talk to Firebase services as before.",
    )
    add_para(
        doc,
        "This document presents the cost-relevant decisions that need a sign-off before "
        "the infrastructure is applied. Three pre-bundled scenarios are offered for "
        "convenience — Lean, Balanced, and Full Production — but each service can also "
        "be sized independently. Every option is reversible with a single configuration "
        "change since the entire stack is held in code (Terraform).",
    )

    # ─── Bundled scenarios up front ──────────────────────────
    add_heading(doc, "Recommended scenarios at a glance", level=1)
    add_para(
        doc,
        "Three bundled scenarios are described below. Pick one as the baseline; "
        "individual services can be adjusted afterwards.",
    )

    headers = ["Scenario", "When to choose", "Est. total /mo"]
    rows = [
        ["Lean (recommended now)", "Pre-launch, no live users, internal QA", "~$63"],
        ["Balanced", "Soft launch with a small group of real users", "~$103"],
        ["Full Production", "Live customers with uptime expectations", "~$170"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)

    add_para(
        doc,
        "Costs above are infrastructure-only and exclude per-request usage (data egress, "
        "Firestore reads, etc.) which are negligible at current traffic.",
        italic=True,
        size=10,
        color=GREY,
    )
    doc.add_paragraph()

    # ─── Section: Cost knobs overview ────────────────────────
    add_heading(doc, "What drives the cost", level=1)
    add_bullet(
        doc,
        "Reserved CPU and RAM run continuously and bill 24/7 — there is no scale-to-zero for the database.",
        bold_lead="Cloud SQL (database)",
    )
    add_bullet(
        doc,
        "Whether to keep a hot-standby replica in a second zone for automatic failover (HA REGIONAL).",
        bold_lead="Cloud SQL availability",
    )
    add_bullet(
        doc,
        "How many warm instances to keep idle. min=0 means cold-starts on first request; min=1+ removes them but bills 24/7.",
        bold_lead="Cloud Run (web app)",
    )
    add_bullet(
        doc,
        "The connector between Cloud Run and the database is billed per-instance per-hour.",
        bold_lead="VPC connector",
    )
    add_bullet(
        doc,
        "Storage cost is small unless very large volumes of video are kept indefinitely. Versioning and retention policies affect this.",
        bold_lead="Cloud Storage",
    )
    doc.add_paragraph()

    # ─── Section: Cloud SQL — production ─────────────────────
    add_heading(doc, "1. Cloud SQL — production database", level=1)
    add_para(
        doc,
        "Production database tier. Staging runs at the small db-g1-small tier separately; "
        "this section is about the production instance only.",
    )

    headers = ["Option", "Tier", "Availability", "Est. $/mo"]
    rows = [
        ["A — Recommended pre-launch", "db-custom-1-3840 (1 vCPU, 3.75 GB)", "ZONAL + PITR", "~$32"],
        ["B — Cheapest",               "db-g1-small (shared CPU, 1.7 GB)", "ZONAL",        "~$17"],
        ["C — Mid-tier with HA",       "db-custom-1-3840 (1 vCPU, 3.75 GB)", "REGIONAL HA + PITR", "~$64"],
        ["D — Full size, no HA",       "db-custom-2-7680 (2 vCPU, 7.5 GB)", "ZONAL + PITR", "~$58"],
        ["E — Full production",        "db-custom-2-7680 (2 vCPU, 7.5 GB)", "REGIONAL HA + PITR", "~$110"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)
    doc.add_paragraph()

    add_option_block(
        doc,
        "Option A — db-custom-1-3840 ZONAL + PITR  (~$32/mo)  · Recommended now",
        "Half the CPU/RAM of the full production tier; no hot standby. PITR is on so any data can still be rolled back in time. Right starting point while there are no live users.",
        [
            "Lowest cost still appropriate for production data.",
            "Same backup and point-in-time-recovery safety net as the larger tier.",
            "Scaling up later is a one-line code change with seconds of downtime; data is preserved.",
        ],
        [
            "Single-zone — a Google zone outage (rare, ~hours per year) takes the database offline.",
            "Smaller CPU/RAM ceiling; heavy queries may run slower.",
        ],
    )
    add_option_block(
        doc,
        "Option E — db-custom-2-7680 REGIONAL HA + PITR  (~$110/mo)  · Full production",
        "Full production-grade configuration. 2 reserved vCPUs, 7.5 GB RAM, hot standby in a second zone with automatic failover.",
        [
            "Highest uptime SLA (99.99%) with automatic cross-zone failover.",
            "Comfortable headroom for traffic and data growth.",
            "No further action when users arrive.",
        ],
        [
            "Highest monthly cost — roughly 3.5× Option A.",
            "Pays the full production bill while there is no production traffic yet.",
        ],
    )
    add_para(
        doc,
        "Options B, C, D are documented in the appendix and remain available "
        "for special cases (cheapest demo, mid-tier soft launch, single-zone large).",
        italic=True,
        size=10,
        color=GREY,
    )

    doc.add_paragraph()

    # ─── Section: Cloud SQL — staging ────────────────────────
    add_heading(doc, "2. Cloud SQL — staging database", level=1)
    add_para(
        doc,
        "Staging is a small, single-zone database used for testing and engineering work. "
        "Cost is small no matter the choice. Default is fine for almost all cases.",
    )
    headers = ["Option", "Tier", "Est. $/mo", "Notes"]
    rows = [
        ["S1 — Recommended", "db-g1-small (shared CPU, 1.7 GB)", "~$17", "Default. Comfortable for QA and engineering."],
        ["S2 — Cheapest",    "db-f1-micro (shared CPU, 0.6 GB)", "~$8",  "Very small; risks running out of memory under realistic queries."],
        ["S3 — Match prod size", "db-custom-1-3840 ZONAL", "~$30", "Useful only if staging tests need realistic performance numbers."],
    ]
    make_table(doc, headers, rows, highlight_idx=0)

    doc.add_paragraph()

    # ─── Section: Cloud Run — production ─────────────────────
    add_heading(doc, "3. Cloud Run — production web app", level=1)
    add_para(
        doc,
        "Cloud Run runs the Next.js web container. Cost depends on idle ('min instances') "
        "and request-driven CPU/memory time. min=0 saves money but adds cold-starts of "
        "1–3 seconds on the first request after idle; min=1+ removes cold-starts but "
        "bills around-the-clock.",
    )
    headers = ["Option", "Min", "Max", "Memory", "Cold-start", "Est. $/mo idle"]
    rows = [
        ["R1 — Lean",     "0", "10",  "512 MB", "Yes (first req)", "~$0–3"],
        ["R2 — Balanced (recommended for soft launch)", "1", "20", "1 GB", "No",  "~$8–15"],
        ["R3 — Full prod","2", "50",  "1 GB",  "No (with redundancy)", "~$15–30"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)
    add_para(
        doc,
        "Request-driven cost on top of the idle baseline is small — Cloud Run free tier "
        "covers the first 2M requests/month. At pre-launch traffic this is effectively zero.",
        italic=True,
        size=10,
        color=GREY,
    )
    doc.add_paragraph()

    # ─── Section: Cloud Run — staging ────────────────────────
    add_heading(doc, "4. Cloud Run — staging web app", level=1)
    headers = ["Option", "Min", "Max", "Memory", "Est. $/mo"]
    rows = [
        ["RS1 — Recommended", "0", "5", "512 MB", "~$0–1"],
        ["RS2 — Always warm",  "1", "5", "512 MB", "~$5"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)
    add_para(
        doc,
        "Cold-starts on staging are acceptable for almost all use cases. Default min=0.",
        italic=True,
        size=10,
        color=GREY,
    )
    doc.add_paragraph()

    # ─── Section: VPC connector ──────────────────────────────
    add_heading(doc, "5. Network — VPC connector", level=1)
    add_para(
        doc,
        "Cloud Run reaches the database through a small managed network bridge "
        "(VPC connector). It is shared between staging and production. The size of the "
        "bridge sets the max throughput between the web app and the database.",
    )
    headers = ["Option", "Machine type", "Instances", "Throughput", "Est. $/mo"]
    rows = [
        ["N1 — Recommended", "e2-micro",      "2–10", "Up to ~200 Mbps", "~$10"],
        ["N2 — Heavy traffic", "e2-standard-4", "2–10", "Up to ~3 Gbps",   "~$80"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)
    add_para(
        doc,
        "N1 is sufficient until the platform sustains hundreds of concurrent active users.",
        italic=True,
        size=10,
        color=GREY,
    )
    doc.add_paragraph()

    # ─── Section: Cloud Storage ──────────────────────────────
    add_heading(doc, "6. Cloud Storage — buckets and retention", level=1)
    add_para(
        doc,
        "Six buckets are provisioned (videos, exports, firebase × staging, prod). The "
        "production buckets keep object versioning on so accidentally overwritten files "
        "can be recovered. The staging buckets delete files older than 30 days to control "
        "cost.",
    )
    headers = ["Option", "Prod versioning", "Staging retention", "Est. $/mo"]
    rows = [
        ["G1 — Recommended", "On",  "Delete > 30 days", "~$2 + traffic"],
        ["G2 — No versioning",  "Off", "Delete > 30 days", "~$1 + traffic"],
        ["G3 — Strict prod retention", "On + 90-day retention lock", "Delete > 30 days", "~$3 + traffic"],
    ]
    make_table(doc, headers, rows, highlight_idx=0)
    add_para(
        doc,
        "Storage is volume-driven. Until significant video volume exists, the difference "
        "between options is a few dollars per month.",
        italic=True,
        size=10,
        color=GREY,
    )
    doc.add_paragraph()

    # ─── Section: Region ────────────────────────────────────
    add_heading(doc, "7. Region", level=1)
    add_para(
        doc,
        "All resources are placed in us-west1 (Oregon) — chosen because the client is a "
        "US-West-based robotics company. All other GCP services (Cloud Run, Cloud SQL, "
        "Firestore, Cloud Storage, Identity Platform) are available there. Switching "
        "regions later is possible but means recreating resources.",
    )
    doc.add_paragraph()

    # ─── Section: Identity, Firestore, Secrets, Scheduler ───
    add_heading(doc, "8. Identity, Firestore, Secrets, Scheduler — fixed", level=1)
    add_para(
        doc,
        "These services are billed primarily by usage and have no meaningful sizing knob. "
        "Defaults are used.",
    )
    add_bullet(doc, "Two tenants in one project. Free tier covers expected user counts.", bold_lead="Identity Platform")
    add_bullet(doc, "Two named databases (staging-db, prod-db). Pay per read/write/storage; free tier substantial.", bold_lead="Firestore")
    add_bullet(doc, "About 18 secrets per environment. ~$0.50/mo total.", bold_lead="Secret Manager")
    add_bullet(doc, "Two cron jobs (one per environment). ~$0.10/mo each.", bold_lead="Cloud Scheduler")
    doc.add_paragraph()

    # ─── Bundled scenarios — full breakdown ─────────────────
    add_heading(doc, "Bundled scenarios — full breakdown", level=1)

    def scenario_rows(label):
        if label == "Lean":
            return [
                ["Cloud SQL prod",   "Option A — db-custom-1-3840 ZONAL + PITR",   "$32"],
                ["Cloud SQL staging","Option S1 — db-g1-small ZONAL",              "$17"],
                ["Cloud Run prod",   "Option R1 — min=0, max=10, 512 MB",          "$0–3"],
                ["Cloud Run staging","Option RS1 — min=0, max=5, 512 MB",          "$0–1"],
                ["VPC connector",    "Option N1 — e2-micro",                        "$10"],
                ["Storage + others", "Default",                                    "$3–5"],
                ["Total (typical idle)", "",                                       "~$63"],
            ]
        if label == "Balanced":
            return [
                ["Cloud SQL prod",   "Option C — db-custom-1-3840 REGIONAL HA + PITR", "$64"],
                ["Cloud SQL staging","Option S1 — db-g1-small ZONAL",                   "$17"],
                ["Cloud Run prod",   "Option R2 — min=1, max=20, 1 GB",                 "$8–15"],
                ["Cloud Run staging","Option RS1 — min=0, max=5, 512 MB",               "$0–1"],
                ["VPC connector",    "Option N1 — e2-micro",                             "$10"],
                ["Storage + others", "Default",                                         "$3–5"],
                ["Total (typical idle)", "",                                            "~$103"],
            ]
        if label == "Full Production":
            return [
                ["Cloud SQL prod",   "Option E — db-custom-2-7680 REGIONAL HA + PITR", "$110"],
                ["Cloud SQL staging","Option S1 — db-g1-small ZONAL",                   "$17"],
                ["Cloud Run prod",   "Option R3 — min=2, max=50, 1 GB",                 "$15–30"],
                ["Cloud Run staging","Option RS1 — min=0, max=5, 512 MB",               "$0–1"],
                ["VPC connector",    "Option N1 — e2-micro (upgrade later if needed)",  "$10"],
                ["Storage + others", "Default",                                         "$3–5"],
                ["Total (typical idle)", "",                                            "~$170"],
            ]
        return []

    for label in ("Lean", "Balanced", "Full Production"):
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        make_table(doc, ["Service", "Option", "Est. $/mo"], scenario_rows(label),
                   highlight_idx=len(scenario_rows(label)) - 1)
        doc.add_paragraph()

    # ─── Recommendation ─────────────────────────────────────
    add_heading(doc, "Recommendation", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Apply the Lean scenario today. ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        "It is the cheapest configuration that is still appropriate for production data, "
        "preserves the PITR safety net, and matches the staging shape for realistic "
        "testing. The platform is pre-launch with no live users; paying for cross-zone "
        "failover and warm instances now is hard to justify before there is anyone to "
        "notice an outage. The whole stack is held in code, so moving to the Balanced or "
        "Full Production scenario is a single change applied with one command on the day "
        "real customer load is expected."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ─── Decision request ───────────────────────────────────
    add_heading(doc, "Decision needed", level=1)
    add_para(
        doc,
        "Confirm one of the three scenarios above, or specify per-service overrides. "
        "Once chosen, the configuration is applied as a single Terraform run (~15-20 "
        "minutes) and produces fully isolated staging and production environments under "
        "the existing GCP project.",
    )

    add_para(
        doc,
        "All choices are reversible. Sizing up or down later is a one-line code change.",
        italic=True,
        size=10,
        color=GREY,
    )

    # ─── Appendix ───────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Appendix — alternative production database options", level=1)

    add_option_block(
        doc,
        "Option B — db-g1-small ZONAL  (~$17/mo)  · Cheapest, demo only",
        "Same instance shape as the staging environment. Shared CPU rather than reserved.",
        [
            "Cheapest possible production database.",
            "Production matches staging exactly.",
        ],
        [
            "Shared CPU means performance varies under load.",
            "Not designed by Google for production data.",
            "Limited memory; risk of slow queries or out-of-memory failures as data grows.",
        ],
    )
    add_option_block(
        doc,
        "Option C — db-custom-1-3840 REGIONAL HA + PITR  (~$64/mo)  · Mid-tier with HA",
        "Smaller CPU/RAM than the full production tier, but keeps the hot-standby replica for automatic failover across zones.",
        [
            "Automatic failover within ~1 minute if a zone fails.",
            "Higher uptime SLA (99.99% vs 99.95% for zonal).",
            "Roughly half the cost of the full production tier while keeping resilience.",
        ],
        [
            "Roughly twice Option A while sharing the same CPU/RAM ceiling.",
            "HA value is small until there are real users who would notice an outage.",
        ],
    )
    add_option_block(
        doc,
        "Option D — db-custom-2-7680 ZONAL + PITR  (~$58/mo)  · Full size, no HA",
        "Full production CPU and RAM, but in a single zone with no hot standby.",
        [
            "Comfortable CPU/RAM headroom for heavier workloads.",
            "Cheaper than full HA configuration.",
        ],
        [
            "Single-zone risk same as Option A.",
            "Pays for headroom that the platform will not use until real users arrive.",
        ],
    )

    out = "/Users/ahmadrafi/Developer/podifi/supervolcano-teleops-prod/docs/product/Infrastructure-Cost-Options.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
