#!/usr/bin/env python3
"""Generate customer-facing Cloud SQL tier options docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def make_table(doc, headers, rows, header_color="1F3864", alt_color="F2F2F2"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], header_color)

    # body
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r % 2 == 1:
                shade_cell(cell, alt_color)
    return table


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SUPERVOLCANO TELEOPS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Cloud SQL Production Tier — Options for Decision")
    run.bold = True
    run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared by Podifi · 2026-05-09 · Confidential — Internal Use Only")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    # Context
    add_heading(doc, "Context", level=1)
    add_para(
        doc,
        "As part of the GCP migration, the Supervolcano platform will run two PostgreSQL "
        "databases under Google Cloud SQL — one for staging and one for production. "
        "Staging is settled at a small instance (db-g1-small, ~$17/mo). The production "
        "instance is the larger cost item and the size you choose now sets the monthly "
        "baseline until the platform is scaled up later.",
    )
    add_para(
        doc,
        "Because the platform is pre-launch with no live users, there is room to choose "
        "a smaller production tier today and scale up via a single configuration change "
        "the moment real users arrive. The options below describe that trade-off in plain "
        "terms. All options remain on Cloud SQL with automated daily backups.",
    )

    doc.add_paragraph()

    # Why prod is more expensive
    add_heading(doc, "What drives Cloud SQL production cost", level=1)
    add_para(
        doc,
        "Three knobs determine the monthly bill: how many CPUs are reserved 24/7, how much "
        "RAM is reserved, and whether a hot-standby replica is run in a second zone for "
        "automatic failover (the 'HA REGIONAL' option). Storage and backups are minor by "
        "comparison.",
    )
    add_bullet(doc, "Reserved CPUs and RAM run continuously — there is no scale-to-zero like the web app.")
    add_bullet(doc, "HA REGIONAL adds a second instance in another zone and roughly doubles the compute portion of the bill.")
    add_bullet(doc, "Point-in-time recovery (PITR) keeps transaction logs for 7 days; a small storage cost.")

    doc.add_paragraph()

    # Options table
    add_heading(doc, "Options", level=1)
    headers = ["Option", "Tier", "Availability", "Est. $/mo", "Best for"]
    rows = [
        ["A — Recommended pre-launch", "db-custom-1-3840 (1 vCPU, 3.75 GB)", "ZONAL + PITR", "~$32", "Pre-launch and quiet pilots"],
        ["B — Cheapest", "db-g1-small (shared CPU, 1.7 GB)", "ZONAL", "~$17", "Demo only — not recommended for prod data"],
        ["C — Mid-tier with HA", "db-custom-1-3840 (1 vCPU, 3.75 GB)", "REGIONAL HA + PITR", "~$64", "Soft launch with low traffic but uptime needed"],
        ["D — Full size, no HA", "db-custom-2-7680 (2 vCPU, 7.5 GB)", "ZONAL + PITR", "~$58", "Heavier workloads where one-zone risk is acceptable"],
        ["E — Full production", "db-custom-2-7680 (2 vCPU, 7.5 GB)", "REGIONAL HA + PITR", "~$110", "Live production with real users and uptime SLA"],
    ]
    make_table(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "All options assume the same database engine (Postgres 16), private networking, "
        "automated daily backups, and the staging instance unchanged."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    # Pros and cons by option
    add_heading(doc, "Pros and cons", level=1)

    def add_option_section(label, intro, pros, cons):
        h = doc.add_paragraph()
        run = h.add_run(label)
        run.bold = True
        run.font.size = Pt(13)
        add_para(doc, intro, italic=False, size=11)
        p = doc.add_paragraph()
        run = p.add_run("Pros")
        run.bold = True
        run.font.size = Pt(11)
        for item in pros:
            add_bullet(doc, item)
        p = doc.add_paragraph()
        run = p.add_run("Cons")
        run.bold = True
        run.font.size = Pt(11)
        for item in cons:
            add_bullet(doc, item)
        doc.add_paragraph()

    add_option_section(
        "Option A — db-custom-1-3840 ZONAL + PITR  (~$32/mo)  · Recommended pre-launch",
        "Half the CPU and RAM of the full production tier; no hot standby. PITR is kept on so any data in production can still be rolled back in time. This is the right starting point while there are no live users.",
        [
            "Lowest cost that is still appropriate for production workloads.",
            "Same backup and PITR safety net as the larger tier.",
            "Scaling up later is a one-line configuration change with seconds of downtime; data is preserved.",
            "Cleanly mirrors the staging shape so test results carry over.",
        ],
        [
            "Single-zone — a Google zone outage (rare, ~hours per year) takes the database offline.",
            "Smaller CPU/RAM ceiling means heavy queries or large jobs may run slower.",
            "Recommended only while pre-launch traffic is low; revisit when real users arrive.",
        ],
    )

    add_option_section(
        "Option B — db-g1-small ZONAL  (~$17/mo)  · Cheapest",
        "Same instance shape as the staging environment. Shared CPU rather than reserved.",
        [
            "Cheapest possible production database.",
            "Simple — production matches staging exactly.",
        ],
        [
            "Shared CPU means performance varies under load and noisy-neighbour effects can occur.",
            "Not designed for production data — Google does not recommend g1-small for live workloads.",
            "Limited memory increases the risk of slow queries or out-of-memory failures as data grows.",
            "Realistically a stop-gap; will need a tier upgrade before any non-trivial customer use.",
        ],
    )

    add_option_section(
        "Option C — db-custom-1-3840 REGIONAL HA + PITR  (~$64/mo)  · Mid-tier with HA",
        "Smaller CPU/RAM than the full production tier, but keeps the hot-standby replica for automatic failover across zones.",
        [
            "Automatic failover within ~1 minute if a zone fails — no manual intervention.",
            "Higher uptime SLA (99.99% vs 99.95% for zonal).",
            "Roughly half the cost of the full production tier while keeping the resilience properties.",
        ],
        [
            "Roughly twice the cost of Option A while sharing the same CPU/RAM ceiling.",
            "HA value is small until there are real users who would notice an outage.",
            "Better fit for soft launch than pre-launch.",
        ],
    )

    add_option_section(
        "Option D — db-custom-2-7680 ZONAL + PITR  (~$58/mo)  · Full size, no HA",
        "Full production CPU and RAM, but in a single zone with no hot standby.",
        [
            "Plenty of CPU/RAM headroom for heavier workloads or larger data volumes.",
            "Cheaper than full HA configuration.",
        ],
        [
            "Single-zone risk same as Option A.",
            "Pays for headroom that the platform will not use until real users arrive.",
            "If reliability matters enough to size up, HA usually matters too — Option E is the natural endpoint.",
        ],
    )

    add_option_section(
        "Option E — db-custom-2-7680 REGIONAL HA + PITR  (~$110/mo)  · Full production",
        "The full production-grade configuration recommended for live workloads.",
        [
            "Highest uptime SLA (99.99%) with automatic cross-zone failover.",
            "2 reserved vCPUs and 7.5 GB RAM — comfortable headroom.",
            "Production-grade today, no further action needed when users arrive.",
        ],
        [
            "Highest monthly cost — roughly 3.5× Option A.",
            "Pays the full production bill while there is no production traffic.",
            "Scaling down later is possible but means brief downtime.",
        ],
    )

    # Recommendation
    add_heading(doc, "Recommendation", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Start with Option A (db-custom-1-3840 ZONAL + PITR, ~$32/mo). "
    )
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        "It is the cheapest configuration that is still appropriate for production data, "
        "preserves the PITR safety net, and matches the staging shape. The platform is "
        "pre-launch with no real users today — paying $110/month for cross-zone failover "
        "before there is anyone to notice an outage is hard to justify. The configuration "
        "is held in code, so the upgrade to Option E is a single line change applied with "
        "one command on the day the first real customer goes live."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Total monthly infrastructure cost under each option")
    run.bold = True
    run.font.size = Pt(12)

    headers = ["Option for prod database", "Cloud SQL prod", "Other GCP*", "Total /mo"]
    rows = [
        ["A — db-custom-1-3840 ZONAL", "$32", "~$35", "~$67"],
        ["B — db-g1-small ZONAL",       "$17", "~$35", "~$52"],
        ["C — db-custom-1-3840 REGIONAL HA", "$64", "~$35", "~$99"],
        ["D — db-custom-2-7680 ZONAL", "$58", "~$35", "~$93"],
        ["E — db-custom-2-7680 REGIONAL HA", "$110", "~$35", "~$145"],
    ]
    make_table(doc, headers, rows)

    p = doc.add_paragraph()
    run = p.add_run(
        "* Other GCP includes Cloud SQL staging (~$17), Cloud Run (~$5), VPC connector "
        "(~$10), Cloud Storage, Secret Manager, and Cloud Scheduler. Excludes data egress "
        "and per-request costs which scale with traffic."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Decision request
    doc.add_paragraph()
    add_heading(doc, "Decision needed", level=1)
    add_para(
        doc,
        "Please confirm which option to apply. Option A is the recommendation; any of the "
        "other options is reasonable and reversible — the configuration lives in code and "
        "scaling up or down later is straightforward.",
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Once selected, the change is made in a single line of infrastructure code "
        "and applied with one command. No data migration is required at this stage as "
        "the database is fresh."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    out = "/Users/ahmadrafi/Developer/podifi/supervolcano-teleops-prod/docs/product/Cloud-SQL-Tier-Options.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
